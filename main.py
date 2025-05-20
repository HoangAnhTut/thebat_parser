import os
import re
import email
from email import policy
from email.parser import BytesParser
from email.header import decode_header
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document

folder_path = "data"
output_data = []

# Чтение и парсинг писем
for filename in os.listdir(folder_path):
    if filename.endswith(".eml"):
        eml_path = os.path.join(folder_path, filename)
        with open(eml_path, "rb") as file:
            msg = BytesParser(policy=policy.default).parse(file)

        date = msg["Date"]
        from_ = msg["From"]
        to = msg["To"]
        subject = msg["Subject"]

        attachment_header = msg["Content-Disposition"]
        body = msg.get_body(preferencelist=("plain", "html"))
        body_content = body.get_content() if body else ""

        # Преобразование HTML в текст
        if body and body.get_content_type() == "text/html":
            soup = BeautifulSoup(body_content, "html.parser")
            body_content = soup.get_text()

        # Очистка от лишних пробелов и пустых строк
        body_content = "\n".join(
            re.sub(r"\s{2,}", " ", line.strip())
            for line in body_content.splitlines()
            if line.strip()
        )

        # Преобразование даты
        try:
            parsed_date = datetime.strptime(date[:-6], "%a, %d %b %Y %H:%M:%S")
            formatted_date = parsed_date.strftime("%d.%m.%Y %H:%M")
        except Exception as e:
            formatted_date = date

        attachments = []
        for part in msg.iter_attachments():
            attach_name = part.get_filename()
            if attach_name:
                attachments.append(attach_name)

        # Сбор данных
        output_data.append({
            "Дата/время": formatted_date,
            "Отправитель (от кого)": from_,
            "Получатель (кому)": to,
            "Содержание письма / Тема": f"Тема: {subject}\n\n{body_content}\n\n",
            "Названия вложений": ", ".join(attachments) if attachments else " "
        })

print(f"\nВсего обработано писем: {len(output_data)}")

# Создание Word-файла и таблицы
doc = Document()
doc.add_heading('Список писем', 0)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Заголовки
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Дата/время'
hdr_cells[1].text = 'Отправитель (от кого)'
hdr_cells[2].text = 'Получатель (кому)'
hdr_cells[3].text = 'Содержание письма / Тема'
hdr_cells[4].text = 'Названия вложений'

# Добавление строк
for i, email_data in enumerate(output_data):
    row_cells = table.add_row().cells
    row_cells[0].text = email_data.get("Дата/время") or ""
    row_cells[1].text = email_data.get("Отправитель (от кого)") or ""
    row_cells[2].text = email_data.get("Получатель (кому)") or ""
    row_cells[3].text = email_data.get("Содержание письма / Тема") or ""
    row_cells[4].text = email_data.get("Названия вложений") or ""

# Сохранение .docx файла
doc_path = "emails.docx"
doc.save(doc_path)

print(f"\nФайл Word с таблицей сохранён: {doc_path}")
print(f"Обработано {len(output_data)} писем.")
